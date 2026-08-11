#!/usr/bin/env python3
"""
ingest.py

Ingests a locally-synced folder of proposal documents (originating from
SharePoint, synced down by you/your team through your normal SharePoint
sync tools -- this script does not talk to SharePoint itself) into a local
SQLite database, following the 5-table schema: proposals, documents,
people, resume_facts, templates.

No AI/ML. Fact extraction here is intentionally simple and rule-based
(regex/keyword matching for certifications, education, employment dates).
It is a starting point meant to be reviewed and corrected, not trusted
blindly -- flag this clearly to reviewers.

EXPECTED FOLDER LAYOUT (matches your SharePoint structure; adjust
DOC_TYPE_BY_FOLDER below if yours differs):

    <root>/
      AGENCY TO# RFP NAME/            <- one folder per proposal/Task Order
        Solicitation/
          rfp.pdf
        Resumes/
          Historical Resumes/         <- past resumes for this or other TOs
            Jane_Doe.docx
          Generated Resumes/          <- resumes this tool has produced
            Jane_Doe_v1.docx

OPTIONAL SHAREPOINT MANIFEST:
If you have a manifest mapping local files to their SharePoint URLs
(e.g. exported from SharePoint or maintained by hand), pass it with
--manifest. Format: a JSON file of {"local_relative_path": "https://..."}.
Without a manifest, sharepoint_url is left NULL and can be filled in later.

USAGE:
    python ingest.py /path/to/local/proposals --db proposals.db
    python ingest.py /path/to/local/proposals --db proposals.db --manifest sharepoint_urls.json
"""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import sqlite3
import sys
from dataclasses import dataclass
from pathlib import Path

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    import docx as docx_lib
except ImportError:
    docx_lib = None


SUPPORTED_TEXT_EXT = {".pdf", ".docx", ".txt"}

SCHEMA = """
CREATE TABLE IF NOT EXISTS proposals (
    id              INTEGER PRIMARY KEY,
    name            TEXT UNIQUE NOT NULL,   -- canonical name, deadline stripped
    folder_name     TEXT,                    -- most recent raw folder name seen (with deadline)
    status          TEXT DEFAULT 'active',
    requirements    TEXT,
    rfp_metadata    TEXT,                    -- JSON: auto-extracted TORFP#, contract type,
                                              -- contracting officer info, etc. -- rule-based,
                                              -- always review before treating as authoritative
    created_at      TEXT DEFAULT (datetime('now')),
    updated_at      TEXT DEFAULT (datetime('now'))
);

CREATE TABLE IF NOT EXISTS people (
    id              INTEGER PRIMARY KEY,
    full_name       TEXT UNIQUE NOT NULL,
    created_at      TEXT DEFAULT (datetime('now'))
);

CREATE TABLE IF NOT EXISTS documents (
    id                 INTEGER PRIMARY KEY,
    proposal_id        INTEGER REFERENCES proposals(id),
    person_id          INTEGER REFERENCES people(id),
    doc_type           TEXT NOT NULL,        -- 'rfp' | 'resume_historical' | 'resume_generated'

    source_system      TEXT DEFAULT 'sharepoint',  -- 'sharepoint' | 'local_only'
    sharepoint_url      TEXT,
    local_cache_path    TEXT UNIQUE NOT NULL,
    cache_synced_at     TEXT,
    cache_stale         INTEGER DEFAULT 0,

    file_ext           TEXT,
    file_size           INTEGER,
    mtime               REAL,
    raw_text            TEXT,
    content_hash         TEXT,               -- sha256 of raw_text, used to detect
                                              -- the exact same resume content reused
                                              -- across multiple proposal folders
    canonical_document_id INTEGER REFERENCES documents(id),
                                              -- set when this document's content is an
                                              -- exact duplicate of an earlier document for
                                              -- the same person -- NULL means this IS the
                                              -- canonical copy that actually owns the
                                              -- extracted resume_facts
    created_at          TEXT DEFAULT (datetime('now')),
    updated_at          TEXT DEFAULT (datetime('now'))
);

CREATE TABLE IF NOT EXISTS resume_facts (
    id                  INTEGER PRIMARY KEY,
    person_id           INTEGER NOT NULL REFERENCES people(id),
    fact_type           TEXT NOT NULL,   -- 'education' | 'certification' |
                                          -- 'project_experience' | 'employment' |
                                          -- 'technology' | 'responsibility' | 'unclassified'
    fact_text           TEXT NOT NULL,
    start_date           TEXT,
    end_date             TEXT,
    source_document_id   INTEGER NOT NULL REFERENCES documents(id),
    source_section        TEXT,
    conflict_flag         INTEGER DEFAULT 0,
    manually_edited       INTEGER DEFAULT 0,
    created_at            TEXT DEFAULT (datetime('now'))
);

CREATE TABLE IF NOT EXISTS templates (
    id                          INTEGER PRIMARY KEY,
    proposal_id                 INTEGER NOT NULL REFERENCES proposals(id),
    version                     INTEGER DEFAULT 1,
    status                      TEXT DEFAULT 'draft',
    page_limit                  INTEGER,
    structure                   TEXT,
    source_resume_document_id   INTEGER REFERENCES documents(id),
    approved_at                 TEXT,
    created_at                  TEXT DEFAULT (datetime('now'))
);

CREATE INDEX IF NOT EXISTS idx_resume_facts_person ON resume_facts(person_id);
CREATE INDEX IF NOT EXISTS idx_resume_facts_type ON resume_facts(person_id, fact_type);
CREATE INDEX IF NOT EXISTS idx_documents_proposal ON documents(proposal_id);
CREATE INDEX IF NOT EXISTS idx_documents_person ON documents(person_id);
CREATE INDEX IF NOT EXISTS idx_templates_proposal ON templates(proposal_id);

CREATE VIRTUAL TABLE IF NOT EXISTS documents_fts USING fts5(
    raw_text, content='documents', content_rowid='id'
);
CREATE TRIGGER IF NOT EXISTS documents_ai AFTER INSERT ON documents BEGIN
    INSERT INTO documents_fts(rowid, raw_text) VALUES (new.id, new.raw_text);
END;
CREATE TRIGGER IF NOT EXISTS documents_ad AFTER DELETE ON documents BEGIN
    INSERT INTO documents_fts(documents_fts, rowid, raw_text) VALUES ('delete', old.id, old.raw_text);
END;
CREATE TRIGGER IF NOT EXISTS documents_au AFTER UPDATE ON documents BEGIN
    INSERT INTO documents_fts(documents_fts, rowid, raw_text) VALUES ('delete', old.id, old.raw_text);
    INSERT INTO documents_fts(rowid, raw_text) VALUES (new.id, new.raw_text);
END;
"""

DOC_TYPE_BY_FOLDER = {
    "Solicitation": "rfp",
    "Resumes/Historical Resumes": "resume_historical",
    "Resumes/Generated Resumes": "resume_generated",
}

# Known certifications to look for -- extend this list as needed.
KNOWN_CERTS = [
    "PMP", "AWS Certified Solutions Architect", "AWS Certified", "CISSP",
    "Security+", "CISA", "CISM", "ITIL", "Scrum Master", "CSM", "PgMP",
    "Six Sigma", "CCNA", "Azure Administrator", "Azure Solutions Architect",
]

DEGREE_KEYWORDS = [
    "Bachelor of", "Master of", "B.S.", "M.S.", "MBA", "Ph.D.", "PhD",
    "Bachelor's", "Master's", "Associate of",
]

DATE_RANGE_RE = re.compile(
    r"(\b(?:19|20)\d{2}\b)\s*[-\u2013\u2014to]{1,4}\s*(\b(?:19|20)\d{2}\b|present|current)",
    re.IGNORECASE,
)

# Maps section header text (many possible real-world variants, since
# different RFPs ask for different header wording) to one canonical
# fact_type. Add new variants here as you encounter them -- this list
# will never cover every possible header, so anything that doesn't
# match falls back to being captured under its own literal header name
# (fact_type='other', source_section=<the header as written>) rather
# than being silently dropped or shredded into unclassified sentences.
HEADER_ALIASES: dict[str, list[str]] = {
    "summary": ["summary", "professional summary", "executive summary", "profile", "overview"],
    "education": ["education", "education & training", "academic background"],
    "certification": ["certifications", "certification", "credentials", "licenses & certifications"],
    "employment": [
        "employment", "employment history", "work history", "professional experience",
        "experience", "relevant experience", "recent experience", "recent work experience",
        "relevant work experience", "project experience", "project history",
        "relevant project experience",
    ],
    "skills": [
        "skills", "technical skills", "core competencies", "key skills", "technologies",
        "core qualifications", "areas of expertise", "qualifications alignment",
        "senior project management qualifications",
    ],
    "years_of_experience": ["years of experience", "total years of experience"],
    "clearance": ["clearance", "security clearance"],
    "training": ["training", "training & development"],
}

# Build a reverse lookup: normalized header text -> canonical fact_type.
_HEADER_LOOKUP: dict[str, str] = {}
for canonical, variants in HEADER_ALIASES.items():
    for variant in variants:
        _HEADER_LOOKUP[variant.lower().strip()] = canonical

HEADER_LINE_RE = re.compile(r"^[A-Za-z][A-Za-z0-9 &/\-']{1,45}$")
BULLET_PREFIX_RE = re.compile(r"^[\u2022\u25cf\-\*\u2013]\s*")


def normalize_header(line: str) -> str | None:
    """Return the canonical fact_type for a header line, if recognized."""
    key = line.lower().strip().rstrip(":")
    return _HEADER_LOOKUP.get(key)


def looks_like_header(line: str, next_line: str | None = None) -> bool:
    """Heuristic: short line, no sentence-ending punctuation, not a bullet,
    consistent with a resume section header.

    Detection has two tiers:
    1. If the line's normalized text exactly matches a known alias in
       HEADER_ALIASES (e.g. 'Certifications', 'certifications', or
       'CERTIFICATIONS'), it's trusted as a header unconditionally --
       casing and what follows it don't matter, since we already know
       this exact wording is a real header from prior data.
    2. Otherwise, fall back to structural guessing for headers we haven't
       seen before: ALL CAPS lines are trusted directly; Title Case lines
       are only treated as headers if the next line is a bullet, since
       otherwise short Title Case sentences/job titles (e.g. 'Task Leader')
       get misdetected.
    """
    stripped = line.strip().rstrip(":")
    if not stripped or len(stripped) > 46:
        return False
    if BULLET_PREFIX_RE.match(line.strip()):
        return False
    if stripped.endswith((".", ",")):
        return False
    if not HEADER_LINE_RE.match(stripped):
        return False

    if normalize_header(stripped) is not None:
        return True

    if stripped.isupper():
        # Single-word ALL CAPS lines under 6 chars are very likely
        # acronym content (PMP, SQL, AWS, CPA, ITIL, CISA) rather than a
        # genuine section header -- real unknown headers we haven't seen
        # before tend to be longer or multi-word (AWARDS, TRAINING,
        # EXPERTISE). This only affects the fallback guess for headers
        # NOT already in HEADER_ALIASES; known short headers like
        # 'SKILLS' are still caught by the exact-alias check above.
        if " " not in stripped and len(stripped) < 6:
            return False
        return True

    if next_line is not None and BULLET_PREFIX_RE.match(next_line.strip()):
        words = stripped.split()
        capitalized = sum(1 for w in words if w[:1].isupper())
        return capitalized >= max(1, len(words) - 1)

    return False

# Strips a trailing "(...)" from a proposal folder name, e.g.
# "BTS TO31 Program Support (Aug 18 3pm)" -> "BTS TO31 Program Support"
# This is what most of your folder names use to hold the current deadline,
# which changes over time -- the canonical proposal name must NOT include it,
# or a deadline update would look like a brand new proposal.
TRAILING_PAREN_RE = re.compile(r"\s*\([^()]*\)\s*$")


def canonicalize_proposal_name(folder_name: str) -> str:
    canonical = TRAILING_PAREN_RE.sub("", folder_name).strip()
    return canonical or folder_name


@dataclass
class FileRecord:
    path: Path
    proposal_name: str
    doc_type: str
    person_name: str | None


def _iter_docx_blocks(document):
    """Yield paragraphs and tables in the actual order they appear in the
    document body. python-docx's .paragraphs and .tables properties each
    return their own type in isolation, with no positional relationship to
    each other -- appending all paragraphs then all tables (the previous
    approach here) silently misattributes table content to whatever header
    happened to be last in the paragraph stream, which is wrong whenever a
    resume's layout uses tables for structure (very common)."""
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def extract_text(path: Path) -> str:
    ext = path.suffix.lower()
    try:
        if ext == ".txt":
            return path.read_text(encoding="utf-8", errors="ignore")
        if ext == ".pdf":
            if pdfplumber is None:
                return ""
            with pdfplumber.open(path) as pdf:
                return "\n".join(page.extract_text() or "" for page in pdf.pages)
        if ext == ".docx":
            if docx_lib is None:
                return ""
            document = docx_lib.Document(str(path))
            parts = []
            for block in _iter_docx_blocks(document):
                if hasattr(block, "text"):  # Paragraph
                    if block.text.strip():
                        parts.append(block.text)
                else:  # Table -- walk cells in reading order (row by row)
                    for row in block.rows:
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if cell_text:
                                parts.append(cell_text)
            return "\n".join(parts)
    except Exception as exc:
        print(f"  [warn] could not extract text from {path}: {exc}", file=sys.stderr)
    return ""


def guess_person_name(filename: str) -> str:
    stem = Path(filename).stem
    stem = re.sub(r"(?i)[\s_-]*(formatted|final|raw|resume|cv|v\d+|draft)+$", "", stem)
    stem = re.sub(r"[_\-]+", " ", stem).strip()
    return stem or filename


def _walk_files(folder: Path):
    if not folder.exists():
        return
    for f in sorted(folder.rglob("*")):
        if not f.is_file() or f.suffix.lower() not in SUPPORTED_TEXT_EXT:
            continue
        if f.name.startswith("~$"):
            # Word/Office temporary lock file, created while the real file
            # is open for editing (e.g. "~$Cedric Schulman.docx" alongside
            # "Cedric Schulman.docx"). Not real content -- always skip.
            continue
        if f.name.startswith("."):
            # Hidden/system files (e.g. macOS .DS_Store companions,
            # sync-client conflict markers) -- also not real content.
            continue
        yield f


def classify_and_collect(root: Path) -> list[FileRecord]:
    records: list[FileRecord] = []
    for proposal_folder in sorted(p for p in root.iterdir() if p.is_dir()):
        proposal_name = proposal_folder.name

        for folder_name, doc_type in DOC_TYPE_BY_FOLDER.items():
            subfolder = proposal_folder / folder_name
            for f in _walk_files(subfolder):
                if doc_type == "rfp":
                    records.append(FileRecord(f, proposal_name, doc_type, None))
                    continue
                # Resumes sometimes sit directly in the Historical/Generated
                # Resumes folder (name comes from the filename), but more
                # often each person has their own subfolder containing
                # possibly-generic filenames like "Resume.docx" -- in that
                # case the real name is the subfolder, not the file. Check
                # both: if the file is nested inside a subfolder here, use
                # that subfolder's name; otherwise fall back to the filename.
                rel_parts = f.relative_to(subfolder).parts
                if len(rel_parts) > 1:
                    person_name = guess_person_name(rel_parts[0])
                else:
                    person_name = guess_person_name(f.name)
                records.append(FileRecord(f, proposal_name, doc_type, person_name))

        for f in proposal_folder.iterdir():
            if f.is_file() and f.suffix.lower() in SUPPORTED_TEXT_EXT:
                records.append(FileRecord(f, proposal_name, "rfp", None))

    return records


_ALL_HEADER_VARIANTS_BY_LENGTH = sorted(
    {v for variants in HEADER_ALIASES.values() for v in variants},
    key=len, reverse=True,
)


def split_inline_headers(lines: list[str]) -> list[str]:
    """Some resume templates style a section header as a bolded run at the
    START of a paragraph, immediately followed by body text with no line
    break (e.g. one single paragraph reading 'CORE QUALIFICATIONS Ms. Doe
    exceeds the requirements for...'). Plain-text extraction can't see the
    bold formatting that visually separates them, so the header and body
    text arrive as one unbroken line and never get detected as a header at
    all -- the whole block then gets misattributed to whichever real
    header came before it.

    This splits such lines into a proper header line + remainder, but only
    when the matched prefix is ALL CAPS in the original text (not just
    case-insensitively similar), to avoid false-splitting ordinary
    sentences that happen to start with a header word in regular sentence
    case (e.g. 'Experience managing federal contracts...' should NOT be
    split just because 'experience' is a known alias).
    """
    result = []
    for line in lines:
        stripped = line.strip()
        if normalize_header(stripped) is not None:
            # Already an exact standalone header on its own line (e.g. a
            # line that IS just "CERTIFICATIONS") -- never attempt to
            # split these; doing so risks a shorter alias variant (e.g.
            # singular "certification") matching a prefix of the word and
            # treating the trailing letters as bogus remainder content.
            result.append(line)
            continue

        matched = None
        for variant in _ALL_HEADER_VARIANTS_BY_LENGTH:
            vlen = len(variant)
            if len(stripped) <= vlen:
                continue
            candidate_prefix = stripped[:vlen]
            if candidate_prefix.lower() != variant.lower():
                continue
            if not candidate_prefix.isupper():
                continue  # only split when styled as a caps header run
            # Require a genuine word boundary right after the matched
            # variant (space/colon/dash/end), not just any character --
            # otherwise "CERTIFICATIONS" (plural) can match the shorter
            # "certification" (singular) variant and treat the trailing
            # "S" as if it were separate content.
            boundary_char = stripped[vlen:vlen + 1]
            if boundary_char and boundary_char.isalnum():
                continue
            remainder = stripped[vlen:].lstrip(" :\t-")
            if remainder:
                matched = (candidate_prefix, remainder)
                break
        if matched:
            result.append(matched[0])
            result.append(matched[1])
        else:
            result.append(line)
    return result


def extract_facts(text: str) -> list[dict]:
    """Structural, header-aware fact extraction.

    Splits the document into sections by detecting header lines, maps each
    header to a canonical fact_type via HEADER_ALIASES (so 'Relevant
    Experience' and 'Recent Experience' land in the same bucket regardless
    of which wording a given RFP's resume format used), and treats each
    bullet/line under a header as one fact.

    Headers that don't match any known alias are NOT dropped -- they're
    captured with fact_type='other' and source_section set to the literal
    header text, so nothing silently disappears and you can see which new
    header variants show up in real resumes and add them to HEADER_ALIASES.

    If no headers are detected at all (pure narrative resume with no
    structure), falls back to the old sentence-based heuristic extraction
    so something is still captured rather than nothing.
    """
    lines = [ln.rstrip() for ln in text.split("\n")]
    lines = [ln for ln in lines if ln.strip()]
    lines = split_inline_headers(lines)

    sections: list[tuple[str, str | None, list[str]]] = []  # (header_text, fact_type_or_None, body_lines)
    current_header = None
    current_fact_type = None
    current_body: list[str] = []
    any_header_found = False

    for i, line in enumerate(lines):
        next_line = lines[i + 1] if i + 1 < len(lines) else None
        if looks_like_header(line, next_line):
            if current_header is not None:
                sections.append((current_header, current_fact_type, current_body))
            current_header = line.strip().rstrip(":")
            current_fact_type = normalize_header(current_header)
            current_body = []
            any_header_found = True
        else:
            current_body.append(line)

    if current_header is not None:
        sections.append((current_header, current_fact_type, current_body))
    elif current_body:
        # no headers ever found -- treat everything as one unheadered block
        sections.append((None, None, current_body))

    if not any_header_found:
        return _extract_facts_sentence_fallback(text)

    facts: list[dict] = []

    for header_text, fact_type, body_lines in sections:
        if header_text is None:
            # content before the first header (or no headers at all) --
            # run the sentence fallback just on this chunk
            facts.extend(_extract_facts_sentence_fallback("\n".join(body_lines)))
            continue

        resolved_type = fact_type or "other"
        source_section = header_text if fact_type is None else None

        if resolved_type == "summary":
            # Summary is prose, not bullets -- keep as one block
            block_text = " ".join(body_lines).strip()
            if block_text:
                facts.append({
                    "fact_type": "summary", "fact_text": block_text,
                    "start_date": None, "end_date": None, "source_section": source_section,
                })
            continue

        for line in body_lines:
            clean_line = BULLET_PREFIX_RE.sub("", line).strip()
            if not clean_line:
                continue
            date_match = DATE_RANGE_RE.search(clean_line)
            facts.append({
                "fact_type": resolved_type,
                "fact_text": clean_line,
                "start_date": date_match.group(1) if date_match else None,
                "end_date": date_match.group(2) if date_match else None,
                "source_section": source_section,
            })

    return facts


def _extract_facts_sentence_fallback(text: str) -> list[dict]:
    """Original sentence-splitting heuristic, used only when no section
    headers are detected at all (pure narrative resume)."""
    facts: list[dict] = []
    lines = [ln.strip() for ln in re.split(r"[\n.]", text) if ln.strip()]

    for line in lines:
        matched = False

        for cert in KNOWN_CERTS:
            if cert.lower() in line.lower():
                facts.append({"fact_type": "certification", "fact_text": line,
                               "start_date": None, "end_date": None, "source_section": None})
                matched = True
                break
        if matched:
            continue

        if any(kw.lower() in line.lower() for kw in DEGREE_KEYWORDS):
            facts.append({"fact_type": "education", "fact_text": line,
                           "start_date": None, "end_date": None, "source_section": None})
            continue

        date_match = DATE_RANGE_RE.search(line)
        if date_match and len(line) > 15:
            facts.append({
                "fact_type": "employment", "fact_text": line,
                "start_date": date_match.group(1), "end_date": date_match.group(2),
                "source_section": None,
            })
            continue

        if len(line) > 40:
            facts.append({"fact_type": "unclassified", "fact_text": line,
                           "start_date": None, "end_date": None, "source_section": None})

    return facts


SOLICITATION_NUMBER_RE = re.compile(
    r"(?:Solicitation\s*(?:No\.?|Number|#)?|TORFP\s*#?|RFQ\s*(?:No\.?|Number|#)?|"
    r"Task\s+Order\s*(?:No\.?|Number|#))\s*[:\-]?\s*([A-Z0-9][A-Z0-9\-]{4,29})",
    re.IGNORECASE,
)

CONTRACT_TYPE_PATTERNS = [
    ("Firm Fixed Price", ["firm fixed price", "ffp"]),
    ("Time and Materials", ["time and materials", "time-and-materials", "t&m"]),
    ("Cost Plus Fixed Fee", ["cost plus fixed fee", "cpff"]),
    ("Cost Reimbursement", ["cost reimbursement", "cost-reimbursement"]),
    ("Hybrid", ["hybrid contract", "hybrid task order", "hybrid (ffp"]),
]

CO_NAME_RE = re.compile(
    r"Contracting\s+Officer\s*(?:\(CO\))?\s*[:\-]?[ \t]*([A-Z][a-zA-Z.'\-]+(?:[ \t]+[A-Z][a-zA-Z.'\-]+){0,3})",
)

EMAIL_RE = re.compile(r"[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}")

PHONE_RE = re.compile(r"(?:\+?1[\s.\-]?)?\(?\d{3}\)?[\s.\-]\d{3}[\s.\-]\d{4}")


def extract_rfp_metadata(text: str) -> dict:
    """Rule-based extraction of standardized RFP metadata fields. These
    fields (solicitation number, contract type, CO contact info) tend to
    follow much more consistent boilerplate patterns across federal
    solicitations than free-form content like 'labor category' or
    'priority topics' did -- but this is still regex/keyword matching,
    not language understanding. Fields that aren't confidently found are
    left null rather than guessed. Always spot-check against the real
    document; treat this as a fast first pass, not a verified result.
    """
    result = {
        "solicitation_number": None,
        "contract_type": None,
        "contracting_officer_name": None,
        "contracting_officer_email": None,
        "contracting_officer_phone": None,
        "_extraction_method": "rule_based_v1",
        "_needs_review": True,
    }

    sol_match = SOLICITATION_NUMBER_RE.search(text)
    if sol_match:
        result["solicitation_number"] = sol_match.group(1).strip()

    text_lower = text.lower()
    for label, keywords in CONTRACT_TYPE_PATTERNS:
        if any(kw in text_lower for kw in keywords):
            result["contract_type"] = label
            break

    co_match = CO_NAME_RE.search(text)
    if co_match:
        result["contracting_officer_name"] = co_match.group(1).strip().split("\n")[0].strip()

    # Look for an email/phone near the words "Contracting Officer" first
    # (more likely to actually belong to the CO, not some other contact
    # elsewhere in the document); fall back to the first one found anywhere.
    co_context = ""
    co_idx = text.lower().find("contracting officer")
    if co_idx != -1:
        co_context = text[co_idx:co_idx + 400]

    email_match = EMAIL_RE.search(co_context) or EMAIL_RE.search(text)
    if email_match:
        result["contracting_officer_email"] = email_match.group(0)

    phone_match = PHONE_RE.search(co_context) or PHONE_RE.search(text)
    if phone_match:
        result["contracting_officer_phone"] = phone_match.group(0)

    return result


def get_or_create(cur: sqlite3.Cursor, table: str, key_col: str, key_val: str) -> int:
    cur.execute(f"SELECT id FROM {table} WHERE {key_col} = ?", (key_val,))
    row = cur.fetchone()
    if row:
        return row[0]
    cur.execute(f"INSERT INTO {table} ({key_col}) VALUES (?)", (key_val,))
    return cur.lastrowid


def get_or_create_proposal(cur: sqlite3.Cursor, raw_folder_name: str) -> int:
    canonical = canonicalize_proposal_name(raw_folder_name)
    cur.execute("SELECT id, folder_name FROM proposals WHERE name = ?", (canonical,))
    row = cur.fetchone()
    if row:
        proposal_id, current_folder_name = row
        if current_folder_name != raw_folder_name:
            # deadline (or other trailing detail) changed since last ingest --
            # update the stored raw folder name but keep the same proposal row
            cur.execute(
                "UPDATE proposals SET folder_name = ?, updated_at = datetime('now') WHERE id = ?",
                (raw_folder_name, proposal_id),
            )
        return proposal_id
    cur.execute(
        "INSERT INTO proposals (name, folder_name) VALUES (?, ?)",
        (canonical, raw_folder_name),
    )
    return cur.lastrowid


def migrate_schema(conn: sqlite3.Connection) -> None:
    """Add columns to existing databases that predate this version of the
    schema, without touching any existing data. Safe to run every time."""
    cur = conn.cursor()
    cur.execute("PRAGMA table_info(resume_facts)")
    existing_cols = {row[1] for row in cur.fetchall()}
    if "manually_edited" not in existing_cols:
        cur.execute("ALTER TABLE resume_facts ADD COLUMN manually_edited INTEGER DEFAULT 0")
        conn.commit()

    cur.execute("PRAGMA table_info(proposals)")
    existing_cols = {row[1] for row in cur.fetchall()}
    if "rfp_metadata" not in existing_cols:
        cur.execute("ALTER TABLE proposals ADD COLUMN rfp_metadata TEXT")
        conn.commit()
    if "rfp_metadata_manually_edited" not in existing_cols:
        cur.execute("ALTER TABLE proposals ADD COLUMN rfp_metadata_manually_edited INTEGER DEFAULT 0")
        conn.commit()

    cur.execute("PRAGMA table_info(documents)")
    existing_cols = {row[1] for row in cur.fetchall()}
    if "content_hash" not in existing_cols:
        cur.execute("ALTER TABLE documents ADD COLUMN content_hash TEXT")
        conn.commit()
    if "canonical_document_id" not in existing_cols:
        cur.execute("ALTER TABLE documents ADD COLUMN canonical_document_id INTEGER REFERENCES documents(id)")
        conn.commit()


def ingest(root: Path, db_path: Path, manifest_path: Path | None, force: bool = False) -> None:
    manifest = {}
    if manifest_path and manifest_path.exists():
        manifest = json.loads(manifest_path.read_text(encoding="utf-8"))

    conn = sqlite3.connect(db_path)
    conn.execute("PRAGMA foreign_keys = ON;")
    cur = conn.cursor()
    cur.executescript(SCHEMA)
    conn.commit()
    migrate_schema(conn)

    records = classify_and_collect(root)
    if not records:
        print(f"No matching files found under {root}. Check your folder layout.")
        return

    inserted, updated, skipped, facts_created = 0, 0, 0, 0

    for rec in records:
        stat = rec.path.stat()
        abs_path = str(rec.path.resolve())
        rel_path = str(rec.path.relative_to(root))

        cur.execute(
            "SELECT id, file_size, mtime FROM documents WHERE local_cache_path = ?", (abs_path,)
        )
        existing = cur.fetchone()

        if existing and not force and existing[1] == stat.st_size and abs(existing[2] - stat.st_mtime) < 1:
            skipped += 1
            continue

        proposal_id = get_or_create_proposal(cur, rec.proposal_name)
        person_id = get_or_create(cur, "people", "full_name", rec.person_name) if rec.person_name else None
        text = extract_text(rec.path)
        sharepoint_url = manifest.get(rel_path)
        content_hash = hashlib.sha256(text.encode("utf-8")).hexdigest() if text.strip() else None

        # Duplicate detection: if this is a resume and its content is byte-
        # identical to another resume already on file for the SAME person
        # (very common -- the same resume file gets copied as-is into many
        # proposal folders), don't extract a second full set of facts.
        # Instead, point this document at the canonical one that already
        # owns the facts, so corrections made in one place apply everywhere
        # this exact resume was reused, and the person's page doesn't show
        # the same content over and over.
        canonical_document_id = None
        if rec.doc_type in ("resume_historical", "resume_generated") and person_id and content_hash:
            cur.execute(
                """SELECT id FROM documents
                   WHERE person_id = ? AND content_hash = ? AND canonical_document_id IS NULL
                   AND local_cache_path != ?
                   ORDER BY id LIMIT 1""",
                (person_id, content_hash, abs_path),
            )
            canonical_row = cur.fetchone()
            if canonical_row:
                canonical_document_id = canonical_row[0]

        if existing:
            doc_id = existing[0]
            cur.execute(
                """UPDATE documents
                   SET proposal_id=?, person_id=?, doc_type=?, sharepoint_url=?,
                       cache_synced_at=datetime('now'), file_ext=?, file_size=?,
                       mtime=?, raw_text=?, content_hash=?, canonical_document_id=?,
                       updated_at=datetime('now')
                   WHERE id=?""",
                (proposal_id, person_id, rec.doc_type, sharepoint_url,
                 rec.path.suffix.lower(), stat.st_size, stat.st_mtime, text,
                 content_hash, canonical_document_id, doc_id),
            )
            cur.execute(
                "DELETE FROM resume_facts WHERE source_document_id = ? AND manually_edited = 0",
                (doc_id,),
            )
            updated += 1
        else:
            cur.execute(
                """INSERT INTO documents
                   (proposal_id, person_id, doc_type, source_system, sharepoint_url,
                    local_cache_path, cache_synced_at, file_ext, file_size, mtime, raw_text,
                    content_hash, canonical_document_id)
                   VALUES (?, ?, ?, 'sharepoint', ?, ?, datetime('now'), ?, ?, ?, ?, ?, ?)""",
                (proposal_id, person_id, rec.doc_type, sharepoint_url, abs_path,
                 rec.path.suffix.lower(), stat.st_size, stat.st_mtime, text,
                 content_hash, canonical_document_id),
            )
            doc_id = cur.lastrowid
            inserted += 1

        if rec.doc_type in ("resume_historical", "resume_generated") and person_id and canonical_document_id is None:
            for fact in extract_facts(text):
                cur.execute(
                    """INSERT INTO resume_facts
                       (person_id, fact_type, fact_text, start_date, end_date,
                        source_document_id, source_section)
                       VALUES (?, ?, ?, ?, ?, ?, ?)""",
                    (person_id, fact["fact_type"], fact["fact_text"],
                     fact["start_date"], fact["end_date"], doc_id, fact.get("source_section")),
                )
                facts_created += 1

        if rec.doc_type == "rfp" and text.strip():
            cur.execute("SELECT rfp_metadata_manually_edited FROM proposals WHERE id = ?", (proposal_id,))
            already_edited = (cur.fetchone() or [0])[0]
            if not already_edited:
                metadata = extract_rfp_metadata(text)
                cur.execute(
                    "UPDATE proposals SET rfp_metadata = ?, updated_at = datetime('now') WHERE id = ?",
                    (json.dumps(metadata, indent=2), proposal_id),
                )

    conn.commit()
    conn.close()

    print(f"Done. Documents inserted: {inserted}, updated: {updated}, skipped: {skipped}")
    print(f"Facts extracted: {facts_created} (rule-based -- review before trusting)")
    print(f"Database: {db_path.resolve()}")
    if not manifest:
        print("[note] No SharePoint manifest provided -- sharepoint_url left NULL for all documents.")


def cleanup_missing_files(conn: sqlite3.Connection) -> int:
    """Mark documents whose local_cache_path no longer exists on disk.
    Does not delete rows (facts/history stay intact) -- just flags them
    so you know the local cache is stale after a folder rename or move.
    """
    cur = conn.cursor()
    cur.execute("SELECT id, local_cache_path FROM documents WHERE cache_stale = 0")
    marked = 0
    for doc_id, path in cur.fetchall():
        if not Path(path).exists():
            cur.execute("UPDATE documents SET cache_stale = 1 WHERE id = ?", (doc_id,))
            marked += 1
    conn.commit()
    return marked


def main():
    parser = argparse.ArgumentParser(
        description="Ingest a locally-synced copy of SharePoint proposal folders into SQLite."
    )
    parser.add_argument("root", type=Path, help="Root folder containing one subfolder per proposal")
    parser.add_argument("--db", type=Path, default=Path("proposals.db"), help="Path to SQLite DB file")
    parser.add_argument(
        "--manifest", type=Path, default=None,
        help="Optional JSON file mapping local relative paths to SharePoint URLs",
    )
    parser.add_argument(
        "--cleanup", action="store_true",
        help="After ingesting, mark any documents whose local_cache_path no longer "
             "exists (e.g. due to a folder rename) as cache_stale=1",
    )
    parser.add_argument(
        "--force", action="store_true",
        help="Re-process every file and re-extract facts even if the file hasn't "
             "changed since the last ingest. Use this after updating extraction logic.",
    )
    args = parser.parse_args()

    if not args.root.exists() or not args.root.is_dir():
        print(f"Error: {args.root} is not a valid directory.", file=sys.stderr)
        sys.exit(1)

    if pdfplumber is None:
        print("[warn] pdfplumber not installed; PDF text will not be extracted.", file=sys.stderr)
    if docx_lib is None:
        print("[warn] python-docx not installed; DOCX text will not be extracted.", file=sys.stderr)

    ingest(args.root, args.db, args.manifest, args.force)

    if args.cleanup:
        conn = sqlite3.connect(args.db)
        marked = cleanup_missing_files(conn)
        conn.close()
        print(f"Cleanup: marked {marked} document(s) as cache_stale (local file no longer found).")


if __name__ == "__main__":
    main()